"""
Routes expected by the existing React dashboard (apiClient + backendClient).

- POST /upload — word estimate (FileInputBar); optional Redis snapshot shares in-memory milestone state across API workers when ``REDIS_URL`` is set; no ``jobs`` DB row until confirm/start
- POST /api/jobs — draft upload (optional flows using apiClient jobs helpers)
- POST /api/jobs/{id}/estimate — word count + INR estimate
- POST /api/jobs/{id}/start — queue same worker as /job/confirm
- GET /api/jobs — job list sidebar
- GET /api/pricing/config — pricing hook
- POST /job/confirm — start translation (background thread)
- GET /job/{job_id} — poll status (milestone DTO)
- GET /api/translation-outputs/{job_id}/translated.docx|.pdf — downloads
"""

from __future__ import annotations

import json
import logging
import math
import re
import threading
import time
import uuid
from datetime import datetime, timezone
from pathlib import Path
from typing import Any

from fastapi import APIRouter, File, HTTPException, Query, Request, UploadFile
from fastapi.responses import FileResponse
from starlette.concurrency import run_in_threadpool
from pydantic import BaseModel, Field
from sqlalchemy import func, or_, select
from sqlalchemy.exc import IntegrityError

from app.core.pipeline_settings import get_pipeline_settings
from app.db.models import DocumentJob, JobStatus, Profile
from app.db.session import get_session_factory
from app.deps.supabase_auth import AuthProfile, resolve_auth_profile_with_anonymous_fallback
from app.services.parser import parse_document
from app.jobs.job_progress import publish_translation_progress, read_translation_progress
from app.jobs.redis_sync import get_sync_redis
from app.services.pipeline_runner import run_pipeline
from app.services.structured_document_builder import structured_json_path_for_docx
from app.services.translation_pdf_export import export_translation_pdf
from app.services.translation_target import (
    download_stem_label,
    normalize_translation_target,
    translation_target_label,
)
from app.utils.translation_output_filenames import translation_output_filename
from app.services.preview_slice import preview_eligibility, truncate_blocks_to_word_budget
from app.payg_pricing import estimate_payg_inr
from app.services.word_credits import (
    add_usage_row,
    apply_word_charge,
    compute_word_charge,
    refresh_subscription_expiry,
)

logger = logging.getLogger(__name__)

api_router = APIRouter(prefix="/api", tags=["legacy-api"])
outputs_router = APIRouter(prefix="/api/translation-outputs", tags=["legacy-outputs"])
plain_router = APIRouter(tags=["legacy"])

_milestone_lock = threading.Lock()
_milestone_jobs: dict[str, "_MilestoneState"] = {}


class _MilestoneState:
    __slots__ = (
        "job_id",
        "input_path",
        "file_name",
        "word_count",
        "estimated_cost",
        "status",
        "progress_percent",
        "error_message",
        "output_docx",
        "output_pdf",
        "created_at",
        "started_at",
        "completed_at",
        "persist_user_id",
        "billing_snapshot",
        "preview_eligible",
        "document_page_count",
        "preview_pages_cap",
        "translation_target",
    )

    def __init__(
        self,
        job_id: str,
        input_path: Path,
        file_name: str,
        created_at: str,
        *,
        word_count: int | None = None,
        estimated_cost: float | None = None,
        persist_user_id: uuid.UUID | None = None,
    ):
        self.job_id = job_id
        self.input_path = input_path
        self.file_name = file_name
        self.created_at = created_at
        self.persist_user_id = persist_user_id
        if word_count is None:
            self.status = "pending_estimate"
            self.word_count = 0
            self.estimated_cost = 0.0
        else:
            self.status = "estimated"
            self.word_count = word_count
            self.estimated_cost = estimated_cost or 0.0
        self.progress_percent = 0
        self.error_message: str | None = None
        self.output_docx: Path | None = None
        self.output_pdf: Path | None = None
        self.started_at: float | None = None
        self.completed_at: float | None = None
        self.billing_snapshot: dict | None = None
        self.preview_eligible: bool = False
        self.document_page_count: int = 1
        self.preview_pages_cap: int = 3
        self.translation_target: str = "hinglish"


_ALLOWED = frozenset({".pdf", ".docx", ".txt", ".epub", ".ebup"})


_LEGACY_MILESTONE_REDIS_SUFFIX = ":milestone_state"
_LEGACY_MILESTONE_REDIS_TTL_SEC = 86400 * 2
_LEGACY_SNAPSHOT_VER = 1


def _legacy_milestone_redis_key(job_id: str) -> str:
    return f"translator:job:{job_id}{_LEGACY_MILESTONE_REDIS_SUFFIX}"


def _legacy_safe_abs_path(path: Path | None) -> str | None:
    if path is None:
        return None
    try:
        return str(path.resolve())
    except OSError:
        return str(path)


def _legacy_milestone_snapshot_to_dict(st: _MilestoneState) -> dict[str, Any]:
    """JSON-serialize milestone UI state for cross-worker restores (needs Redis)."""
    oux = _legacy_safe_abs_path(st.output_docx)
    updf = _legacy_safe_abs_path(st.output_pdf)
    persist = (
        str(st.persist_user_id) if st.persist_user_id is not None else None
    )
    bs = st.billing_snapshot
    bs_out: dict | None = bs if bs is None or isinstance(bs, dict) else None
    return {
        "v": _LEGACY_SNAPSHOT_VER,
        "job_id": st.job_id,
        "input_relpath": st.input_path.name,
        "file_name": st.file_name,
        "created_at": st.created_at,
        "word_count": max(0, int(st.word_count)),
        "estimated_cost": float(st.estimated_cost or 0),
        "status": st.status,
        "progress_percent": int(st.progress_percent or 0),
        "error_message": st.error_message,
        "output_docx_abs": oux,
        "output_pdf_abs": updf,
        "persist_user_id": persist,
        "billing_snapshot": bs_out,
        "preview_eligible": bool(st.preview_eligible),
        "document_page_count": int(st.document_page_count or 1),
        "preview_pages_cap": int(st.preview_pages_cap or 3),
        "translation_target": str(st.translation_target or "hinglish"),
    }


def _legacy_milestone_from_snapshot_payload(
    d: dict[str, Any], *, settings: Any
) -> _MilestoneState | None:
    if int(d.get("v") or 0) != _LEGACY_SNAPSHOT_VER:
        return None
    job_id_val = str(d.get("job_id") or "").strip()
    rel = str(d.get("input_relpath") or "").strip()
    fn = str(d.get("file_name") or "").strip()
    created = str(d.get("created_at") or "")
    uid_raw = d.get("persist_user_id")
    try:
        persist: uuid.UUID | None = (
            uuid.UUID(str(uid_raw)) if uid_raw else None
        )
    except (ValueError, TypeError):
        persist = None
    doc = settings.temp_dir / rel
    if not doc.is_file():
        logger.debug(
            "skip milestone Redis restore — missing upload file (%s)",
            doc,
        )
        return None
    status_saved = str(d.get("status") or "estimated")

    wc = max(0, int(d.get("word_count") or 0))
    ecost = float(d.get("estimated_cost") or 0)

    if status_saved == "pending_estimate":
        st = _MilestoneState(
            job_id_val,
            doc,
            fn or "upload",
            created or datetime.now(timezone.utc).isoformat(),
            persist_user_id=persist,
        )
        st.preview_eligible = bool(d.get("preview_eligible", False))
        st.document_page_count = max(1, int(d.get("document_page_count") or 1))
        st.preview_pages_cap = max(1, int(d.get("preview_pages_cap") or 3))
        st.translation_target = str(
            d.get("translation_target") or st.translation_target
        )
        st.progress_percent = int(d.get("progress_percent") or 0)
    else:
        st = _MilestoneState(
            job_id_val,
            doc,
            fn or "upload",
            created or datetime.now(timezone.utc).isoformat(),
            word_count=wc,
            estimated_cost=ecost,
            persist_user_id=persist,
        )
        st.status = status_saved
        st.progress_percent = int(d.get("progress_percent") or 0)
        st.preview_eligible = bool(d.get("preview_eligible", False))
        st.document_page_count = max(1, int(d.get("document_page_count") or 1))
        st.preview_pages_cap = max(1, int(d.get("preview_pages_cap") or 3))
        st.translation_target = str(
            d.get("translation_target") or st.translation_target
        )

    em = d.get("error_message")
    st.error_message = str(em) if em is not None else None

    bs = d.get("billing_snapshot")
    st.billing_snapshot = bs if isinstance(bs, dict) else None

    od = d.get("output_docx_abs")
    pd = d.get("output_pdf_abs")

    op_docx = Path(str(od)) if od else None
    if op_docx is not None:
        if not op_docx.is_file():
            op_docx = None
        st.output_docx = op_docx
    op_pdf = Path(str(pd)) if pd else None
    if op_pdf is not None:
        if not op_pdf.is_file():
            op_pdf = None
        st.output_pdf = op_pdf

    return st


def _legacy_maybe_restore_milestone_from_redis(job_id: str) -> None:
    """If milestone state exists in Redis but not in this worker, reconstruct it locally."""
    settings = get_pipeline_settings()
    if not settings.redis_url:
        return
    with _milestone_lock:
        if job_id in _milestone_jobs:
            return
    try:
        r = get_sync_redis()
        raw = r.get(_legacy_milestone_redis_key(job_id))
        if not raw:
            return
        payload = json.loads(raw)
        if not isinstance(payload, dict):
            return
    except Exception:
        logger.debug("milestone Redis restore parse failed job_id=%s", job_id, exc_info=True)
        return
    st = _legacy_milestone_from_snapshot_payload(payload, settings=settings)
    if st is None:
        return
    with _milestone_lock:
        if job_id in _milestone_jobs:
            return
        _milestone_jobs[job_id] = st
    logger.info("Restored milestone job state from Redis (job_id=%s)", job_id)


def _legacy_publish_milestone(job_id: str) -> None:
    """Persist current in-process milestone snapshot to Redis (multi-worker deployments)."""
    settings = get_pipeline_settings()
    if not settings.redis_url:
        return
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
        if st is None:
            return
        payload = _legacy_milestone_snapshot_to_dict(st)
    try:
        get_sync_redis().set(
            _legacy_milestone_redis_key(job_id),
            json.dumps(payload, separators=(",", ":")),
            ex=_LEGACY_MILESTONE_REDIS_TTL_SEC,
        )
    except Exception:
        logger.debug(
            "milestone Redis publish failed job_id=%s",
            job_id,
            exc_info=True,
        )


def _milestone_ui_status_to_db_job_status(milestone_status: str) -> str:
    """Map in-memory milestone labels to ``public.jobs.status`` (JobStatus values)."""
    m = (milestone_status or "").lower()
    if m == "completed":
        return JobStatus.COMPLETED.value
    if m == "preview_ready":
        return JobStatus.PREVIEW_READY.value
    if m == "failed":
        return JobStatus.FAILED.value
    if m == "processing":
        return JobStatus.PROCESSING.value
    if m == "awaiting_payment":
        return JobStatus.AWAITING_PAYMENT.value
    return JobStatus.PENDING.value


def _legacy_try_profile(request: Request) -> AuthProfile | None:
    """Optional auth for dashboard routes: persists to ``public.jobs`` when DB + user context exist."""
    settings = get_pipeline_settings()
    auth = request.headers.get("Authorization")
    key = request.headers.get("X-API-Key")
    can_anon = bool(
        settings.allow_anonymous_jobs and settings.anonymous_job_user_id
    )
    if not auth and not key:
        if can_anon:
            return resolve_auth_profile_with_anonymous_fallback(None, None)
        return None
    return resolve_auth_profile_with_anonymous_fallback(auth, key)


def _legacy_sync_milestone_to_db_on_commit(
    job_id: str, profile: AuthProfile
) -> None:
    """Create or update ``jobs`` row when the user starts translation (confirm/start).

    Estimates alone do not write to the DB — avoids table growth from word-count-only use.
    """
    factory = get_session_factory()
    if factory is None:
        return
    try:
        jid = uuid.UUID(job_id)
    except ValueError:
        return
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
        if st is None:
            return
        fn = st.file_name
        rel = st.input_path.name
        suffix = Path(fn).suffix.lower().lstrip(".") or None
        status = _milestone_ui_status_to_db_job_status(st.status)
        wc = max(0, int(st.word_count))
        cost = float(st.estimated_cost or 0)
        quote_inr = (
            round(cost, 2) if (st.status or "").lower() == "awaiting_payment" else 0.0
        )
        tt = normalize_translation_target(getattr(st, "translation_target", None))
    try:
        with factory() as session:
            row = session.get(DocumentJob, jid)
            if row is None:
                session.add(
                    DocumentJob(
                        id=jid,
                        user_id=profile.id,
                        status=status,
                        input_filename=fn,
                        input_file_path=rel,
                        export_format="docx",
                        file_type=suffix,
                        tokens_used=wc,
                        cost_inr=cost,
                        quoted_payg_inr=quote_inr,
                        translation_target=tt,
                    )
                )
            else:
                row.status = status
                row.input_filename = fn
                row.input_file_path = rel
                row.file_type = suffix
                row.tokens_used = wc
                row.cost_inr = cost
                row.quoted_payg_inr = quote_inr
                row.translation_target = tt
            session.commit()
        logger.info("Synced committed job %s to public.jobs", job_id)
    except IntegrityError:
        logger.warning(
            "Could not insert job %s (missing profiles row for user?)",
            job_id,
            exc_info=True,
        )
    except Exception:
        logger.exception("Could not sync job %s to database", job_id)


def _legacy_update_job_row(
    job_id_str: str,
    *,
    status: str,
    output_docx: Path | None = None,
    error_message: str | None = None,
    tokens_used: int | None = None,
    cost_inr: float | None = None,
) -> None:
    factory = get_session_factory()
    if factory is None:
        return
    try:
        jid = uuid.UUID(job_id_str)
    except ValueError:
        return
    try:
        with factory() as session:
            row = session.get(DocumentJob, jid)
            if row is None:
                uid: uuid.UUID | None = None
                fn = "upload"
                rel: str | None = None
                ft: str | None = None
                tok = tokens_used
                cost = cost_inr
                with _milestone_lock:
                    st = _milestone_jobs.get(job_id_str)
                    if st is not None and st.persist_user_id is not None:
                        uid = st.persist_user_id
                        fn = st.file_name or fn
                        try:
                            rel = st.input_path.name
                        except Exception:
                            rel = None
                        sfx = Path(fn).suffix.lower().lstrip(".")
                        ft = sfx or None
                        if tok is None:
                            tok = max(0, int(st.word_count))
                        if cost is None:
                            cost = float(st.estimated_cost or 0)
                if uid is None:
                    logger.warning(
                        "Legacy job %s: cannot persist status %s (no jobs row and no user on milestone).",
                        job_id_str,
                        status,
                    )
                    return
                tt_ins = "hinglish"
                with _milestone_lock:
                    st_m = _milestone_jobs.get(job_id_str)
                    if st_m is not None:
                        tt_ins = normalize_translation_target(
                            getattr(st_m, "translation_target", None)
                        )
                row = DocumentJob(
                    id=jid,
                    user_id=uid,
                    status=status,
                    input_filename=fn,
                    input_file_path=rel,
                    export_format="docx",
                    file_type=ft,
                    tokens_used=tok if tok is not None else 0,
                    cost_inr=cost if cost is not None else 0,
                    translation_target=tt_ins,
                )
                session.add(row)
            row.status = status
            if output_docx is not None:
                row.output_file_path = str(output_docx)
            if error_message is not None:
                row.error_message = error_message[:8000]
            if tokens_used is not None:
                row.tokens_used = max(0, int(tokens_used))
            if cost_inr is not None:
                row.cost_inr = cost_inr
            if status == JobStatus.COMPLETED.value:
                row.completed_at = datetime.now(timezone.utc)
            session.commit()
    except Exception:
        logger.exception("Could not update legacy job %s in database", job_id_str)


def _estimate_word_count_fast(path: Path) -> int:
    """
    Fast word count for pricing / UI (plain extraction).
    Avoids full structured parse — that path can take minutes on large PDFs.
    """
    suffix = path.suffix.lower()
    if suffix == ".txt":
        raw = path.read_text(encoding="utf-8", errors="replace")
        return len(re.findall(r"\S+", raw))
    if suffix == ".pdf":
        import fitz

        doc = fitz.open(path)
        try:
            chunks: list[str] = []
            for page in doc:
                chunks.append(page.get_text())
        finally:
            doc.close()
        return len(re.findall(r"\S+", " ".join(chunks)))
    if suffix == ".docx":
        from docx import Document

        doc = Document(path)
        parts: list[str] = []
        for p in doc.paragraphs:
            if p.text and p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    t = (cell.text or "").strip()
                    if t:
                        parts.append(t)
        return len(re.findall(r"\S+", " ".join(parts)))
    if suffix in {".epub", ".ebup"}:
        from app.services.parser.epub_parser import parse_epub

        blocks = parse_epub(path)
        parts: list[str] = []
        for b in blocks:
            if b.text:
                parts.append(b.text)
            if b.data:
                for row in b.data:
                    parts.extend(cell for cell in row if cell)
        return len(re.findall(r"\S+", " ".join(parts)))
    raise ValueError(f"Unsupported type: {suffix}")


def _count_words(path: Path) -> int:
    blocks = parse_document(path)
    parts: list[str] = []
    for b in blocks:
        if b.text:
            parts.append(b.text)
        if b.data:
            for row in b.data:
                parts.extend(cell for cell in row if cell)
    return len(re.findall(r"\S+", " ".join(parts)))


def _empty_billing_dict(total_words: int) -> dict:
    return {
        "total_words": max(0, int(total_words)),
        "free_used": 0,
        "subscription_used": 0,
        "remaining_words": 0,
        "amount_to_pay": float(estimate_payg_inr(max(0, int(total_words)))),
        "user_plan_type": "payg",
    }


def _billing_for_profile_word_count(profile_id: uuid.UUID, word_count: int) -> dict:
    factory = get_session_factory()
    if factory is None:
        return _empty_billing_dict(word_count)
    try:
        with factory() as session:
            p = session.get(Profile, profile_id)
            if p is None:
                return _empty_billing_dict(word_count)
            refresh_subscription_expiry(p)
            session.commit()
            b = compute_word_charge(p, word_count)
            return b.as_dict()
    except Exception:
        logger.exception("billing preview failed")
        return _empty_billing_dict(word_count)


def _assert_wallet_covers_payg_slice(profile_id: uuid.UUID, payg_charge_inr: float) -> None:
    """Ensure INR wallet can cover the PAYG slice; does not recompute word charge.

    Use the same ``amount_to_pay`` as ``_billing_for_profile_word_count`` for this request
    (a second :func:`compute_word_charge` in the same request could diverge).
    """
    payg = float(payg_charge_inr or 0)
    if payg <= 1e-9:
        return
    if not get_pipeline_settings().payg_checkout_required:
        return
    factory = get_session_factory()
    if factory is None:
        return
    with factory() as session:
        p = session.get(Profile, profile_id)
        if p is None:
            raise HTTPException(status_code=404, detail="Profile not found.")
        bal = float(p.credits_inr_balance or 0)
    if payg > bal + 1e-9:
        raise HTTPException(
            status_code=402,
            detail=(
                f"Pay-as-you-go for this job is ₹{payg:.2f}; your account shows ₹{bal:.2f}. "
                f"Use Razorpay checkout on the job page, or for local dev set "
                f"PAYG_CHECKOUT_REQUIRED=false (or legacy PAYG_WALLET_REQUIRED=false)."
            ),
        )


def _legacy_settle_milestone_billing(job_id: str) -> None:
    factory = get_session_factory()
    if factory is None:
        return
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
        if st is None or st.persist_user_id is None:
            return
        uid = st.persist_user_id
        wc = max(0, int(st.word_count))
    try:
        jid = uuid.UUID(job_id)
    except ValueError:
        return
    snap: dict | None = None
    try:
        with factory() as session:
            p = session.get(Profile, uid)
            if p is None:
                return
            refresh_subscription_expiry(p)
            b = compute_word_charge(p, wc)
            apply_word_charge(session, p, b)
            add_usage_row(
                session,
                user_id=uid,
                job_id=jid,
                word_units=b.total_words,
                payg_inr=b.amount_to_pay,
            )
            session.commit()
            snap = b.as_dict()
    except Exception:
        logger.exception("Could not settle milestone billing for job %s", job_id)
        return
    if snap is None:
        return
    with _milestone_lock:
        st2 = _milestone_jobs.get(job_id)
        if st2 is not None:
            st2.billing_snapshot = snap
            st2.estimated_cost = float(snap.get("amount_to_pay") or 0)


def _run_milestone_job(job_id: str) -> None:
    settings = get_pipeline_settings()
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
        if st is None:
            return
        st.status = "processing"
        st.progress_percent = 10
        st.started_at = time.monotonic()
        path = st.input_path
        tt = normalize_translation_target(getattr(st, "translation_target", None))

    _legacy_publish_milestone(job_id)
    _legacy_update_job_row(job_id, status=JobStatus.PROCESSING.value)

    if not settings.openai_api_key:
        with _milestone_lock:
            st = _milestone_jobs.get(job_id)
            if st:
                st.status = "failed"
                st.error_message = "OPENAI_API_KEY is not configured."
                st.progress_percent = 0
        _legacy_update_job_row(
            job_id,
            status=JobStatus.FAILED.value,
            error_message="OPENAI_API_KEY is not configured.",
        )
        _legacy_publish_milestone(job_id)
        return

    try:

        def report_progress(pct: int) -> None:
            p = max(0, min(100, int(pct)))
            with _milestone_lock:
                st = _milestone_jobs.get(job_id)
                if st:
                    st.progress_percent = max(int(st.progress_percent or 0), p)

        docx = run_pipeline(
            path,
            on_progress=report_progress,
            progress_job_id=job_id,
            translation_target=tt,
        )
        publish_translation_progress(
            job_id,
            progress_percent=97,
            current_stage="converting_to_pdf",
        )
        with _milestone_lock:
            st = _milestone_jobs.get(job_id)
            if st:
                st.progress_percent = 97
        pdf: Path | None = None
        try:
            pdf = export_translation_pdf(docx, structured_json_path_for_docx(docx))
        except Exception as e:
            logger.warning(
                "PDF export failed for job %s: %s",
                job_id,
                e,
                exc_info=logger.isEnabledFor(logging.DEBUG),
            )
            pdf = None
        with _milestone_lock:
            st = _milestone_jobs.get(job_id)
            if st:
                st.output_docx = docx
                st.output_pdf = pdf
                st.status = "completed"
                st.progress_percent = 100
                st.completed_at = time.monotonic()
        _legacy_publish_milestone(job_id)
        _legacy_settle_milestone_billing(job_id)
        wc, ecost = 0, 0.0
        with _milestone_lock:
            st = _milestone_jobs.get(job_id)
            if st:
                wc = st.word_count
                ecost = float(st.estimated_cost or 0)
        _legacy_update_job_row(
            job_id,
            status=JobStatus.COMPLETED.value,
            output_docx=docx,
            tokens_used=wc,
            cost_inr=ecost,
        )
        _legacy_publish_milestone(job_id)
    except Exception as e:
        with _milestone_lock:
            st = _milestone_jobs.get(job_id)
            if st:
                st.status = "failed"
                st.error_message = str(e)[:2000]
                st.progress_percent = 0
        err = str(e)[:2000]
        _legacy_update_job_row(
            job_id,
            status=JobStatus.FAILED.value,
            error_message=err,
        )
        _legacy_publish_milestone(job_id)


def activate_milestone_after_payg_payment(job_id_str: str) -> None:
    """After Razorpay verifies a per-job PAYG order, move milestone out of awaiting_payment and run the pipeline."""
    _legacy_maybe_restore_milestone_from_redis(job_id_str)
    with _milestone_lock:
        st = _milestone_jobs.get(job_id_str)
        if st is None:
            logger.warning(
                "activate_milestone_after_payg_payment: no milestone state for job_id=%s",
                job_id_str,
            )
            return
        if st.status != "awaiting_payment":
            return
        st.status = "queued"
        st.progress_percent = 5
    _legacy_publish_milestone(job_id_str)
    t = threading.Thread(target=_run_milestone_job, args=(job_id_str,), daemon=True)
    t.start()


def _consume_preview_quota(profile_id: uuid.UUID) -> None:
    factory = get_session_factory()
    if factory is None:
        return
    today = datetime.now(timezone.utc).date()
    max_per = max(1, int(get_pipeline_settings().translation_preview_max_starts_per_day))
    try:
        with factory() as session:
            p = session.get(Profile, profile_id)
            if p is None:
                raise HTTPException(status_code=404, detail="Profile not found.")
            if p.preview_quota_utc_date != today:
                p.preview_quota_utc_date = today
                p.preview_quota_count = 0
            if int(p.preview_quota_count or 0) >= max_per:
                raise HTTPException(
                    status_code=429,
                    detail=(
                        f"Free preview limit reached ({max_per} per UTC day). "
                        "Try again tomorrow or translate the full document."
                    ),
                )
            p.preview_quota_count = int(p.preview_quota_count or 0) + 1
            session.commit()
    except HTTPException:
        raise
    except Exception as e:
        logger.exception("preview quota update failed")
        raise HTTPException(
            status_code=503,
            detail="Could not verify preview quota.",
        ) from e


def _run_milestone_preview_job(job_id: str) -> None:
    settings = get_pipeline_settings()
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
        if st is None:
            return
        st.status = "processing"
        st.progress_percent = 10
        st.started_at = time.monotonic()
        st.completed_at = None
        path = st.input_path
        cap = max(1, int(st.preview_pages_cap))
        tt = normalize_translation_target(getattr(st, "translation_target", None))

    _legacy_publish_milestone(job_id)
    _legacy_update_job_row(job_id, status=JobStatus.PROCESSING.value)

    if not settings.openai_api_key:
        with _milestone_lock:
            st = _milestone_jobs.get(job_id)
            if st:
                st.status = "failed"
                st.error_message = "OPENAI_API_KEY is not configured."
                st.progress_percent = 0
        _legacy_update_job_row(
            job_id,
            status=JobStatus.FAILED.value,
            error_message="OPENAI_API_KEY is not configured.",
        )
        _legacy_publish_milestone(job_id)
        return

    try:

        def report_progress(pct: int, *, stage: str = "chunk_translating") -> None:
            p = max(0, min(100, int(pct)))
            with _milestone_lock:
                st = _milestone_jobs.get(job_id)
                if st:
                    st.progress_percent = max(int(st.progress_percent or 0), p)
                    p = int(st.progress_percent)
            publish_translation_progress(
                job_id,
                progress_percent=p,
                current_stage=stage,
            )

        # Bounded parse (first N PDF pages or word budget for DOCX/TXT) keeps preview fast.
        parse_stop = threading.Event()

        def _parse_pulse() -> None:
            x = 10
            while not parse_stop.wait(timeout=1.5):
                x = min(13, x + 1)
                report_progress(x, stage="parsing_document")

        pulse_th = threading.Thread(target=_parse_pulse, daemon=True)
        pulse_th.start()
        try:
            wpp = max(50, int(settings.translation_preview_words_per_page_estimate))
            word_ceiling = min(
                cap * wpp,
                max(150, int(settings.translation_preview_max_words)),
            )
            sfx = path.suffix.lower()
            if sfx == ".pdf":
                all_blocks = parse_document(path, max_pdf_pages=cap)
            else:
                all_blocks = parse_document(path, max_preview_words=word_ceiling)
            preview_blocks = truncate_blocks_to_word_budget(
                all_blocks,
                max_words=word_ceiling,
            )
        finally:
            parse_stop.set()
            pulse_th.join(timeout=0.5)

        report_progress(14, stage="parsing_document")
        if not preview_blocks:
            raise ValueError("No translatable content in the preview range.")

        docx = run_pipeline(
            path,
            blocks=preview_blocks,
            on_progress=report_progress,
            progress_job_id=job_id,
            translation_target=tt,
        )
        pdf: Path | None = None
        try:
            pdf = export_translation_pdf(docx, structured_json_path_for_docx(docx))
        except Exception as e:
            logger.warning(
                "PDF export failed for preview job %s: %s",
                job_id,
                e,
                exc_info=logger.isEnabledFor(logging.DEBUG),
            )
            pdf = None
        with _milestone_lock:
            st = _milestone_jobs.get(job_id)
            if st:
                st.output_docx = docx
                st.output_pdf = pdf
                st.status = "preview_ready"
                st.progress_percent = 100
                st.completed_at = time.monotonic()
        _legacy_publish_milestone(job_id)
        _legacy_update_job_row(
            job_id,
            status=JobStatus.PREVIEW_READY.value,
            output_docx=docx,
        )
        _legacy_publish_milestone(job_id)
    except Exception as e:
        with _milestone_lock:
            st = _milestone_jobs.get(job_id)
            if st:
                st.status = "failed"
                st.error_message = str(e)[:2000]
                st.progress_percent = 0
        err = str(e)[:2000]
        _legacy_update_job_row(
            job_id,
            status=JobStatus.FAILED.value,
            error_message=err,
        )
        _legacy_publish_milestone(job_id)


class UploadEstimateResponse(BaseModel):
    job_id: str
    file_name: str
    word_count: int
    estimated_cost: float
    total_words: int = 0
    free_used: int = 0
    subscription_used: int = 0
    remaining_words: int = 0
    amount_to_pay: float = 0
    user_plan_type: str = "payg"
    preview_eligible: bool = False
    document_page_count: int = 1
    preview_pages_cap: int = 3


class JobConfirmBody(BaseModel):
    job_id: str
    input_lang: str = "en"
    translation_target: str | None = None


class JobDetailDto(BaseModel):
    job_id: str
    status: str
    progress_percent: int
    word_count: int
    estimated_cost: float
    error_message: str | None = None
    current_stage: str | None = None
    batches_done: int | None = None
    batches_total: int | None = None
    segments_translated: int | None = None
    segments_total: int | None = None
    output_docx_available: bool = False
    output_pdf_available: bool = False
    translation_duration_seconds: float | None = None
    output_pdf_hint: str | None = Field(
        default=None,
        description="Shown when the job finished but PDF export is unavailable on the server.",
    )
    total_words: int | None = None
    free_used: int | None = None
    subscription_used: int | None = None
    remaining_words: int | None = None
    amount_to_pay: float | None = None
    user_plan_type: str | None = None
    preview_mode: bool = False
    preview_eligible: bool | None = None
    document_page_count: int | None = None
    preview_pages_cap: int | None = None
    translation_target: str | None = None
    translation_target_label: str | None = None


def _translation_target_from_confirm_body(body: JobConfirmBody) -> str:
    if body.translation_target and str(body.translation_target).strip():
        return normalize_translation_target(body.translation_target)
    return normalize_translation_target("hinglish")


@plain_router.post("/upload", response_model=UploadEstimateResponse)
async def legacy_upload(request: Request, file: UploadFile = File(...)):
    settings = get_pipeline_settings()
    profile = _legacy_try_profile(request)
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename required.")
    suffix = Path(file.filename).suffix.lower()
    if suffix not in _ALLOWED:
        raise HTTPException(
            status_code=415,
            detail="Only PDF, DOCX, EPUB, and TXT are supported.",
        )
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file.")
    if settings.max_upload_bytes > 0 and len(data) > settings.max_upload_bytes:
        mb = round(settings.max_upload_bytes / (1024 * 1024), 1)
        raise HTTPException(
            status_code=413,
            detail=f"File too large (max {mb} MB on API). Raise MAX_UPLOAD_BYTES to increase.",
        )

    job_id = str(uuid.uuid4())
    path = settings.temp_dir / f"milestone_{job_id}{suffix}"
    path.write_bytes(data)
    try:
        # Avoid blocking the event loop on large PDFs / slow disks (keeps /health + proxy alive).
        wc = await run_in_threadpool(_estimate_word_count_fast, path)
    except Exception as e:
        path.unlink(missing_ok=True)
        raise HTTPException(
            status_code=422,
            detail=f"Could not read document: {e}",
        ) from e

    cost = estimate_payg_inr(wc)
    bill = _empty_billing_dict(wc)
    if profile:
        bill = _billing_for_profile_word_count(profile.id, wc)
    pay_amount = float(bill.get("amount_to_pay", cost))

    created = datetime.now(timezone.utc).isoformat()
    pe, dpages, pcap = preview_eligibility(path, wc)
    with _milestone_lock:
        st = _MilestoneState(
            job_id,
            path,
            file.filename or "upload",
            created,
            word_count=wc,
            estimated_cost=pay_amount,
            persist_user_id=profile.id if profile else None,
        )
        st.preview_eligible = pe
        st.document_page_count = dpages
        st.preview_pages_cap = pcap
        _milestone_jobs[job_id] = st

    _legacy_publish_milestone(job_id)
    return UploadEstimateResponse(
        job_id=job_id,
        file_name=file.filename or "upload",
        word_count=wc,
        estimated_cost=pay_amount,
        total_words=int(bill["total_words"]),
        free_used=int(bill["free_used"]),
        subscription_used=int(bill["subscription_used"]),
        remaining_words=int(bill["remaining_words"]),
        amount_to_pay=float(bill["amount_to_pay"]),
        user_plan_type=str(bill["user_plan_type"]),
        preview_eligible=pe,
        document_page_count=dpages,
        preview_pages_cap=pcap,
    )


@api_router.get("/pricing/config")
def legacy_pricing_config():
    settings = get_pipeline_settings()
    from app.billing_constants import FREE_CREDITS_INITIAL
    from app.services.word_credits import SUBSCRIPTION_CREDITS_MONTHLY

    return {
        "free_credits_words": FREE_CREDITS_INITIAL,
        "subscription_inr_monthly": 999,
        "subscription_inr_yearly": 9990,
        "subscription_words_per_cycle": SUBSCRIPTION_CREDITS_MONTHLY,
        "rate_inr_per_10000_words": float(settings.rate_inr_per_10000_words),
        "minimum_charge_inr": 5.0,
        "payg_checkout_required": bool(settings.payg_checkout_required),
        "max_upload_file_mb": (
            settings.max_upload_bytes // (1024 * 1024)
            if settings.max_upload_bytes > 0
            else 0
        ),
        "currency": "INR",
        "trust_payment_copy": (
            "Payments are processed over encrypted connections. "
            "We never store full card numbers."
        ),
        "free_tier_label": "Free credits",
        "payg_label": "Pay-as-you-go",
        "subscription_label": "Monthly",
        "translation_preview_max_pages": int(settings.translation_preview_max_pages),
        "translation_preview_max_starts_per_day": int(
            settings.translation_preview_max_starts_per_day
        ),
        "translation_preview_max_words": int(settings.translation_preview_max_words),
    }


@api_router.post("/jobs")
async def legacy_api_create_job(request: Request, file: UploadFile = File(...)):
    """UploadPage: create draft job; estimate comes from POST /api/jobs/{id}/estimate."""
    settings = get_pipeline_settings()
    profile = _legacy_try_profile(request)
    if not file.filename:
        raise HTTPException(status_code=400, detail="Filename required.")
    suffix = Path(file.filename).suffix.lower()
    if suffix not in _ALLOWED:
        raise HTTPException(
            status_code=415,
            detail="Only PDF, DOCX, EPUB, and TXT are supported.",
        )
    data = await file.read()
    if not data:
        raise HTTPException(status_code=400, detail="Empty file.")
    if settings.max_upload_bytes > 0 and len(data) > settings.max_upload_bytes:
        mb = round(settings.max_upload_bytes / (1024 * 1024), 1)
        raise HTTPException(
            status_code=413,
            detail=f"File too large (max {mb} MB on API). Raise MAX_UPLOAD_BYTES to increase.",
        )

    job_id = str(uuid.uuid4())
    path = settings.temp_dir / f"milestone_{job_id}{suffix}"
    path.write_bytes(data)

    created = datetime.now(timezone.utc).isoformat()
    with _milestone_lock:
        _milestone_jobs[job_id] = _MilestoneState(
            job_id,
            path,
            file.filename or "upload",
            created,
            persist_user_id=profile.id if profile else None,
        )
    _legacy_publish_milestone(job_id)
    return {"id": job_id}


@api_router.post("/jobs/{job_id}/estimate")
def legacy_api_estimate(job_id: str):
    _legacy_maybe_restore_milestone_from_redis(job_id)
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
        if st is None:
            raise HTTPException(status_code=404, detail="Job not found.")
        if st.status != "pending_estimate":
            early = {
                "jobId": st.job_id,
                "wordCount": st.word_count,
                "amountCents": int(round(st.estimated_cost * 100)),
                "currency": "INR",
            }
        else:
            early = None
            input_path = st.input_path
    if early is not None:
        _legacy_publish_milestone(job_id)
        return early

    try:
        wc = _estimate_word_count_fast(input_path)
    except Exception as e:
        raise HTTPException(
            status_code=422,
            detail=f"Could not read document: {e}",
        ) from e
    cost = estimate_payg_inr(wc)

    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
        if st is None:
            raise HTTPException(status_code=404, detail="Job not found.")
        if st.status == "pending_estimate":
            st.word_count = wc
            st.estimated_cost = cost
            st.status = "estimated"
            pe, dpages, pcap = preview_eligibility(input_path, wc)
            st.preview_eligible = pe
            st.document_page_count = dpages
            st.preview_pages_cap = pcap

        snapshot = {
            "jobId": st.job_id,
            "wordCount": st.word_count,
            "amountCents": int(round(st.estimated_cost * 100)),
            "currency": "INR",
        }

    _legacy_publish_milestone(job_id)
    return {
        "jobId": snapshot["jobId"],
        "wordCount": snapshot["wordCount"],
        "amountCents": snapshot["amountCents"],
        "currency": snapshot["currency"],
    }


@api_router.post("/jobs/{job_id}/start")
def legacy_api_start(request: Request, job_id: str):
    profile = _legacy_try_profile(request)
    _legacy_maybe_restore_milestone_from_redis(job_id)
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
        if st is None:
            raise HTTPException(status_code=404, detail="Job not found.")
        if st.status != "estimated":
            raise HTTPException(
                status_code=409,
                detail="Job must be estimated before start.",
            )
        if profile:
            st.persist_user_id = profile.id
        st.status = "queued"
        st.progress_percent = 5

    if profile:
        _legacy_sync_milestone_to_db_on_commit(job_id, profile)

    _legacy_publish_milestone(job_id)
    t = threading.Thread(target=_run_milestone_job, args=(job_id,), daemon=True)
    t.start()
    return {"status": "queued"}


def _document_job_to_list_dict(job: DocumentJob) -> dict:
    created = job.created_at.isoformat() if job.created_at else ""
    cost = float(job.cost_inr) if job.cost_inr is not None else 0.0
    out: dict = {
        "id": str(job.id),
        "fileName": job.input_filename or "upload",
        "status": job.status,
        "createdAt": created,
        "amountCents": int(round(cost * 100)),
        "currency": "INR",
    }
    if job.completed_at is not None:
        out["completedAt"] = job.completed_at.isoformat()
    return out


def _resolve_document_job_output_docx(row: DocumentJob) -> Path | None:
    """Locate translated DOCX from DB row (absolute path or relative to ``data_dir``).

    ``output_file_path`` may point at the primary deliverable only: ``.docx``, or when the
    user chose PDF/ZIP export, ``.pdf`` / ``.zip``. The DOCX is always written next to those
    files with the same stem (see ``rq_tasks`` finalize). Resolving via ``.with_suffix`` fixes
    DOCX downloads when the DB row references PDF or ZIP.
    """
    if not row.output_file_path:
        return None
    raw = str(row.output_file_path).strip()
    if not raw:
        return None
    settings = get_pipeline_settings()
    primary = Path(raw)
    if not primary.is_absolute():
        primary = (settings.data_dir / raw).resolve()
    else:
        primary = primary.resolve()

    if not primary.is_file():
        return None

    suf = primary.suffix.lower()
    if suf == ".docx":
        return primary
    if suf in (".pdf", ".zip"):
        docx = primary.with_suffix(".docx")
        return docx if docx.is_file() else None
    return None


def _apply_redis_progress_fields(
    rp: dict,
 *,
    progress_percent: int,
    stage: str,
) -> tuple[
    int,
    str,
    int | None,
    int | None,
    int | None,
    int | None,
]:
    """Normalize ``read_translation_progress`` payload into DTO fields."""
    progress_percent = int(rp.get("progress_percent", progress_percent))
    stage = str(rp.get("current_stage") or stage)
    batches_done = rp.get("batches_done")
    batches_total = rp.get("batches_total")
    segments_translated = rp.get("segments_translated")
    segments_total = rp.get("segments_total")
    if isinstance(batches_done, (int, float)):
        batches_done = int(batches_done)
    else:
        batches_done = None
    if isinstance(batches_total, (int, float)):
        batches_total = int(batches_total)
    else:
        batches_total = None
    if isinstance(segments_translated, (int, float)):
        segments_translated = int(segments_translated)
    else:
        segments_translated = None
    if isinstance(segments_total, (int, float)):
        segments_total = int(segments_total)
    else:
        segments_total = None
    return (
        progress_percent,
        stage,
        batches_done,
        batches_total,
        segments_translated,
        segments_total,
    )


def _document_job_row_to_job_detail_dto(row: DocumentJob) -> JobDetailDto:
    """Build milestone poll DTO from a persisted ``jobs`` row (server restart / RQ worker)."""
    st = (row.status or "").strip().lower()
    progress_percent = 0
    stage: str | None = None
    batches_done: int | None = None
    batches_total: int | None = None
    segments_translated: int | None = None
    segments_total: int | None = None
    tt = normalize_translation_target(getattr(row, "translation_target", None))
    ttl = translation_target_label(tt)

    if st == JobStatus.COMPLETED.value:
        progress_percent = 100
    elif st == JobStatus.PREVIEW_READY.value:
        progress_percent = 100
    elif st == JobStatus.FAILED.value:
        progress_percent = 0
    elif st == JobStatus.PROCESSING.value:
        # Low default when Redis has no key yet — avoids a fake "stuck at 35%" (prepare
        # can take minutes before first publish, or if REDIS_URL is unset).
        stage = "preparing"
        progress_percent = 12
        rp = read_translation_progress(str(row.id))
        if rp:
            (
                progress_percent,
                stage,
                batches_done,
                batches_total,
                segments_translated,
                segments_total,
            ) = _apply_redis_progress_fields(
                rp, progress_percent=progress_percent, stage=stage
            )
            rttl = rp.get("translation_target_label")
            if rttl:
                ttl = str(rttl)
    else:
        stage = "queued"
        progress_percent = 10

    docx_path = _resolve_document_job_output_docx(row)
    docx_ok = bool(docx_path and docx_path.is_file())
    pdf_ok = False
    pdf_hint: str | None = None
    if docx_ok:
        pdf_path = docx_path.with_suffix(".pdf")
        pdf_ok = pdf_path.is_file()
    if (
        st in (JobStatus.COMPLETED.value, JobStatus.PREVIEW_READY.value)
        and docx_ok
        and not pdf_ok
    ):
        pdf_hint = (
            "PDF export failed after translation (WeasyPrint from structured content, then "
            "DOCX→PDF fallback). Check server logs: WeasyPrint needs Pango/Cairo/GTK on the host; "
            "fallback needs LibreOffice, ReportLab, or Word. DOCX download still works."
        )

    dur: float | None = None
    if row.processing_time_seconds is not None:
        dur = float(row.processing_time_seconds)

    wc = max(0, int(row.tokens_used or 0))
    cost = float(row.cost_inr or 0)
    err = row.error_message if st == JobStatus.FAILED.value else None

    bill_live: dict | None = None
    if st == JobStatus.PREVIEW_READY.value:
        bill_live = _billing_for_profile_word_count(row.user_id, wc)

    return JobDetailDto(
        job_id=str(row.id),
        status=st,
        progress_percent=max(0, min(100, int(progress_percent))),
        word_count=wc,
        estimated_cost=cost,
        error_message=err,
        current_stage=stage,
        batches_done=batches_done,
        batches_total=batches_total,
        segments_translated=segments_translated,
        segments_total=segments_total,
        output_docx_available=docx_ok,
        output_pdf_available=pdf_ok,
        translation_duration_seconds=dur,
        output_pdf_hint=pdf_hint,
        total_words=int(bill_live["total_words"]) if bill_live else None,
        free_used=int(bill_live["free_used"]) if bill_live else None,
        subscription_used=int(bill_live["subscription_used"]) if bill_live else None,
        remaining_words=int(bill_live["remaining_words"]) if bill_live else None,
        amount_to_pay=float(bill_live["amount_to_pay"]) if bill_live else None,
        user_plan_type=str(bill_live["user_plan_type"]) if bill_live else None,
        preview_mode=st == JobStatus.PREVIEW_READY.value,
        preview_eligible=None,
        document_page_count=None,
        preview_pages_cap=None,
        translation_target=tt,
        translation_target_label=ttl,
    )


@api_router.get("/jobs")
def legacy_list_jobs(
    request: Request,
    completed_only: bool = Query(
        False,
        description="If true, return only completed rows from public.jobs (paginated JSON object).",
    ),
    page: int = Query(1, ge=1),
    page_size: int = Query(10, ge=1, le=50),
):
    """List jobs from Supabase/Postgres for the signed-in user, overlaid with in-process milestone state.

    When ``completed_only`` is true, the response is a JSON object
    ``{ items, total, page, pageSize }`` (no in-memory merge). Otherwise returns a JSON array (legacy).
    """
    profile = _legacy_try_profile(request)
    factory = get_session_factory()

    if completed_only:
        if not get_pipeline_settings().database_url:
            raise HTTPException(
                status_code=503,
                detail=(
                    "Database not configured. Set SUPABASE_DATABASE_URL (or DATABASE_URL) to your "
                    "Supabase Transaction pooler URI so completed jobs can be read from public.jobs."
                ),
            )
        if not profile:
            raise HTTPException(
                status_code=401,
                detail="Sign in to load translation history.",
            )
        if not factory:
            raise HTTPException(
                status_code=503,
                detail="Database session unavailable. Check SUPABASE_DATABASE_URL and network access.",
            )
        try:
            with factory() as session:
                done_clause = or_(
                    DocumentJob.status == JobStatus.COMPLETED.value,
                    DocumentJob.completed_at.isnot(None),
                )
                filt = (DocumentJob.user_id == profile.id, done_clause)
                total_raw = session.scalar(
                    select(func.count()).select_from(DocumentJob).where(*filt)
                )
                total = int(total_raw or 0)
                offset = (page - 1) * page_size
                stmt = (
                    select(DocumentJob)
                    .where(*filt)
                    .order_by(DocumentJob.created_at.desc())
                    .offset(offset)
                    .limit(page_size)
                )
                items = [
                    _document_job_to_list_dict(j) for j in session.scalars(stmt)
                ]
            return {
                "items": items,
                "total": total,
                "page": page,
                "pageSize": page_size,
            }
        except HTTPException:
            raise
        except Exception as e:
            logger.exception("legacy_list_jobs completed_only: database query failed")
            raise HTTPException(
                status_code=503,
                detail=(
                    "Could not load history from the database. Confirm Supabase has public.jobs "
                    "and the API can reach SUPABASE_DATABASE_URL."
                ),
            ) from e

    by_id: dict[str, dict] = {}

    if profile and factory:
        try:
            with factory() as session:
                stmt = (
                    select(DocumentJob)
                    .where(DocumentJob.user_id == profile.id)
                    .order_by(DocumentJob.created_at.desc())
                    .limit(500)
                )
                for job in session.scalars(stmt):
                    by_id[str(job.id)] = _document_job_to_list_dict(job)
        except Exception:
            logger.exception("legacy_list_jobs: database query failed")

    with _milestone_lock:
        for st in _milestone_jobs.values():
            if profile is not None:
                # Keep jobs tied to this user, plus unscoped in-memory jobs (persist is None),
                # e.g. upload before auth resolved or DB insert failed — otherwise History stays empty.
                if (
                    st.persist_user_id is not None
                    and st.persist_user_id != profile.id
                ):
                    continue
            else:
                if st.persist_user_id is not None:
                    continue
            by_id[st.job_id] = {
                "id": st.job_id,
                "fileName": st.file_name,
                "status": st.status,
                "createdAt": st.created_at,
                "amountCents": int(round(st.estimated_cost * 100)),
                "currency": "INR",
            }

    rows = list(by_id.values())
    rows.sort(key=lambda r: r["createdAt"], reverse=True)
    return rows


@plain_router.post("/job/confirm")
def legacy_job_confirm(request: Request, body: JobConfirmBody):
    profile = _legacy_try_profile(request)
    jid = body.job_id
    tgt = _translation_target_from_confirm_body(body)
    _legacy_maybe_restore_milestone_from_redis(jid)
    with _milestone_lock:
        st = _milestone_jobs.get(jid)
        if st is None:
            raise HTTPException(status_code=404, detail="Job not found.")
        if st.status not in ("estimated", "preview_ready"):
            raise HTTPException(status_code=409, detail="Job already started or completed.")
        st.translation_target = tgt
        wc = max(0, int(st.word_count))
        if profile:
            st.persist_user_id = profile.id
        if st.status == "preview_ready":
            st.output_docx = None
            st.output_pdf = None
            st.completed_at = None

    breakdown: dict
    if profile and get_session_factory():
        breakdown = _billing_for_profile_word_count(profile.id, wc)
    else:
        breakdown = _empty_billing_dict(wc)
    payg = float(breakdown.get("amount_to_pay") or 0)
    settings = get_pipeline_settings()
    requires_razorpay_checkout = payg > 1e-9 and bool(settings.payg_checkout_required)

    if requires_razorpay_checkout:
        if profile is None:
            raise HTTPException(
                status_code=401,
                detail="Sign in to pay for this translation.",
            )
        with _milestone_lock:
            st2 = _milestone_jobs.get(jid)
            if st2 is None:
                raise HTTPException(status_code=404, detail="Job not found.")
            st2.translation_target = tgt
            st2.status = "awaiting_payment"
            st2.progress_percent = 0
        _legacy_sync_milestone_to_db_on_commit(jid, profile)
        base: dict = {
            "ok": True,
            "awaiting_payment": True,
            "amount_to_pay": payg,
        }
        base.update(breakdown)
        _legacy_publish_milestone(jid)
        return base

    if profile and get_session_factory():
        _assert_wallet_covers_payg_slice(profile.id, payg)

    with _milestone_lock:
        st3 = _milestone_jobs.get(jid)
        if st3 is None:
            raise HTTPException(status_code=404, detail="Job not found.")
        st3.translation_target = tgt
        st3.status = "queued"
        st3.progress_percent = 5

    if profile:
        _legacy_sync_milestone_to_db_on_commit(jid, profile)

    _legacy_publish_milestone(jid)
    t = threading.Thread(target=_run_milestone_job, args=(jid,), daemon=True)
    t.start()
    base = {"ok": True, "awaiting_payment": False}
    base.update(breakdown)
    return base


@plain_router.post("/job/preview-start")
def legacy_preview_start(request: Request, body: JobConfirmBody):
    """Translate only the first ``preview_pages_cap`` pages (free; quota + sign-in required)."""
    profile = _legacy_try_profile(request)
    if profile is None:
        raise HTTPException(
            status_code=401,
            detail="Sign in to start a free preview.",
        )
    jid = body.job_id
    tgt = _translation_target_from_confirm_body(body)
    _legacy_maybe_restore_milestone_from_redis(jid)
    with _milestone_lock:
        st = _milestone_jobs.get(jid)
        if st is None:
            raise HTTPException(status_code=404, detail="Job not found.")
        if st.status != "estimated":
            raise HTTPException(
                status_code=409,
                detail="Preview can only start from an estimated job.",
            )
        if not st.preview_eligible:
            raise HTTPException(
                status_code=400,
                detail=(
                    "This document is too short for a free preview "
                    f"(about {st.document_page_count} page(s); need more than "
                    f"{st.preview_pages_cap}). Use full translation instead."
                ),
            )
        st.translation_target = tgt
        st.persist_user_id = profile.id

    _consume_preview_quota(profile.id)

    with _milestone_lock:
        st2 = _milestone_jobs.get(jid)
        if st2 is None:
            raise HTTPException(status_code=404, detail="Job not found.")
        st2.status = "queued"
        st2.progress_percent = 5

    _legacy_sync_milestone_to_db_on_commit(jid, profile)
    _legacy_publish_milestone(jid)
    t = threading.Thread(target=_run_milestone_preview_job, args=(jid,), daemon=True)
    t.start()
    return {"ok": True, "status": "queued"}


@plain_router.get("/job/{job_id}", response_model=JobDetailDto)
def legacy_job_status(request: Request, job_id: str):
    _legacy_maybe_restore_milestone_from_redis(job_id)
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
    if st is not None:
        dur: float | None = None
        if st.started_at is not None and st.completed_at is not None:
            dur = st.completed_at - st.started_at

        progress_percent = int(st.progress_percent or 0)
        stage = None
        batches_done: int | None = None
        batches_total: int | None = None
        segments_translated: int | None = None
        segments_total: int | None = None
        if st.status == "processing":
            stage = "chunk_translating"
            rp = read_translation_progress(job_id)
            if rp:
                (
                    rp_pct,
                    stage,
                    batches_done,
                    batches_total,
                    segments_translated,
                    segments_total,
                ) = _apply_redis_progress_fields(
                    rp,
                    progress_percent=progress_percent,
                    stage=stage,
                )
                progress_percent = max(progress_percent, rp_pct)
            elif progress_percent < 14:
                stage = "parsing_document"

        docx_ok = st.output_docx is not None and st.output_docx.is_file()
        pdf_ok = st.output_pdf is not None and st.output_pdf.is_file()
        pdf_hint: str | None = None
        if (
            st.status in ("completed", "preview_ready")
            and docx_ok
            and not pdf_ok
        ):
            pdf_hint = (
                "PDF export failed after translation (WeasyPrint from structured content, then "
                "DOCX→PDF fallback). Check server logs: WeasyPrint needs Pango/Cairo/GTK on the host; "
                "fallback needs LibreOffice, ReportLab, or Word. DOCX download still works."
            )

        snap = st.billing_snapshot
        bill_live: dict | None = None
        if snap is None and st.persist_user_id is not None and st.word_count >= 0:
            bill_live = _billing_for_profile_word_count(
                st.persist_user_id, st.word_count
            )
        tw = int(snap["total_words"]) if snap else (
            int(bill_live["total_words"]) if bill_live else None
        )
        fu = int(snap["free_used"]) if snap else (
            int(bill_live["free_used"]) if bill_live else None
        )
        su = int(snap["subscription_used"]) if snap else (
            int(bill_live["subscription_used"]) if bill_live else None
        )
        rw = int(snap["remaining_words"]) if snap else (
            int(bill_live["remaining_words"]) if bill_live else None
        )
        atp = float(snap["amount_to_pay"]) if snap else (
            float(bill_live["amount_to_pay"]) if bill_live else None
        )
        upt = str(snap["user_plan_type"]) if snap else (
            str(bill_live["user_plan_type"]) if bill_live else None
        )
        tt = normalize_translation_target(getattr(st, "translation_target", None))
        ttl = translation_target_label(tt)
        if st.status == "processing":
            rp2 = read_translation_progress(job_id)
            if rp2 and rp2.get("translation_target_label"):
                ttl = str(rp2["translation_target_label"])
        return JobDetailDto(
            job_id=st.job_id,
            status=st.status,
            progress_percent=progress_percent,
            word_count=st.word_count,
            estimated_cost=st.estimated_cost,
            error_message=st.error_message,
            current_stage=stage,
            batches_done=batches_done,
            batches_total=batches_total,
            segments_translated=segments_translated,
            segments_total=segments_total,
            output_docx_available=docx_ok,
            output_pdf_available=pdf_ok,
            translation_duration_seconds=dur,
            output_pdf_hint=pdf_hint,
            total_words=tw,
            free_used=fu,
            subscription_used=su,
            remaining_words=rw,
            amount_to_pay=atp,
            user_plan_type=upt,
            preview_mode=st.status == "preview_ready",
            preview_eligible=st.preview_eligible if st.status == "estimated" else None,
            document_page_count=st.document_page_count,
            preview_pages_cap=st.preview_pages_cap,
            translation_target=tt,
            translation_target_label=ttl,
        )

    factory = get_session_factory()
    if factory is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    try:
        jid = uuid.UUID(job_id)
    except ValueError as e:
        raise HTTPException(status_code=400, detail="Invalid job id.") from e

    profile = _legacy_try_profile(request)
    with factory() as session:
        row = session.get(DocumentJob, jid)
    if row is None:
        raise HTTPException(status_code=404, detail="Job not found.")
    if profile is None:
        raise HTTPException(
            status_code=401,
            detail="Sign in to view this job.",
        )
    if row.user_id != profile.id:
        raise HTTPException(status_code=403, detail="Not your job.")

    return _document_job_row_to_job_detail_dto(row)


def _legacy_translated_download_filename(job_id: str, ext: str, fallback: Path) -> str:
    """Attachment name ``{upload_stem}-Hinglish.{ext}``; falls back to ``fallback.name``."""
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
        if st is not None and (st.file_name or "").strip():
            ol = download_stem_label(getattr(st, "translation_target", None))
            return translation_output_filename(
                st.file_name, ext, output_language=ol
            )
    factory = get_session_factory()
    if factory is not None:
        try:
            jid = uuid.UUID(job_id)
        except ValueError:
            return fallback.name
        with factory() as session:
            row = session.get(DocumentJob, jid)
        if row is not None and (row.input_filename or "").strip():
            ol = download_stem_label(getattr(row, "translation_target", None))
            return translation_output_filename(
                row.input_filename, ext, output_language=ol
            )
    return fallback.name


def _legacy_output_docx_path_for_user(job_id: str, request: Request) -> Path:
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
    if st is not None and st.output_docx is not None and st.output_docx.is_file():
        return st.output_docx

    factory = get_session_factory()
    if factory is None:
        raise HTTPException(status_code=503, detail="Database is not configured.")
    try:
        jid = uuid.UUID(job_id)
    except ValueError as e:
        raise HTTPException(status_code=400, detail="Invalid job id.") from e

    profile = _legacy_try_profile(request)
    if profile is None:
        raise HTTPException(status_code=401, detail="Sign in to download.")

    with factory() as session:
        row = session.get(DocumentJob, jid)
    if row is None:
        raise HTTPException(status_code=404, detail="DOCX not ready.")
    if row.user_id != profile.id:
        raise HTTPException(status_code=403, detail="Not your job.")
    if row.status not in (
        JobStatus.COMPLETED.value,
        JobStatus.PREVIEW_READY.value,
    ):
        raise HTTPException(status_code=404, detail="DOCX not ready.")

    path = _resolve_document_job_output_docx(row)
    if path is None or not path.is_file():
        raise HTTPException(status_code=404, detail="DOCX not ready.")
    return path


@outputs_router.get("/{job_id}/translated.docx")
def legacy_download_docx(request: Request, job_id: str):
    path = _legacy_output_docx_path_for_user(job_id, request)
    return FileResponse(
        path=path,
        filename=_legacy_translated_download_filename(job_id, "docx", path),
        media_type=(
            "application/vnd.openxmlformats-officedocument."
            "wordprocessingml.document"
        ),
    )


@outputs_router.get("/{job_id}/translated.pdf")
def legacy_download_pdf(request: Request, job_id: str):
    _legacy_maybe_restore_milestone_from_redis(job_id)
    with _milestone_lock:
        st = _milestone_jobs.get(job_id)
    if st is not None and st.output_pdf is not None and st.output_pdf.is_file():
        return FileResponse(
            path=st.output_pdf,
            filename=_legacy_translated_download_filename(job_id, "pdf", st.output_pdf),
            media_type="application/pdf",
        )

    docx_path = _legacy_output_docx_path_for_user(job_id, request)
    pdf_path = docx_path.with_suffix(".pdf")
    if not pdf_path.is_file():
        raise HTTPException(status_code=404, detail="PDF not available.")
    return FileResponse(
        path=pdf_path,
        filename=_legacy_translated_download_filename(job_id, "pdf", pdf_path),
        media_type="application/pdf",
    )
