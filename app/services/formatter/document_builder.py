"""Rebuild a translated DOCX from classified blocks (book-style Libre Baskerville typography)."""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt

from app.models.document_models import (
    BlockType,
    ClassifiedBlock,
    ContentBlock,
    SectionAction,
    StructuralTag,
)
from app.services.formatter.book_typography import (
    LIBRE_BASKERVILLE,
    strip_markdown_artifacts,
)
from app.services.formatter.docx_embed_fonts import embed_libre_baskerville

STYLE_BOOK_TITLE = "Translator Book Title"
STYLE_BOOK_AUTHOR = "Translator Book Author"
STYLE_TOC_HEADING = "Translator TOC Heading"
STYLE_TOC_ENTRY = "Translator TOC Entry"


def iter_auto_generated_toc_entries(
    _body_blocks: list[ClassifiedBlock],
) -> list[ClassifiedBlock]:
    """
    Reserved for a future auto-generated TOC from document headings.

    Returns an empty list today; callers may extend ``toc_items`` with this result
    once heading extraction and page mapping are implemented.
    """
    return []


def _clear_paragraph_runs(p) -> None:
    for child in list(p._element):
        if child.tag == qn("w:r"):
            p._element.remove(child)


def _add_multiline_run(p, text: str, *, bold: bool, size_pt: float) -> None:
    _clear_paragraph_runs(p)
    run = p.add_run()
    parts = (text or "").split("\n")
    run.text = parts[0] if parts else ""
    _set_run_fonts(run, bold=bold)
    run.font.size = Pt(size_pt)
    for line in parts[1:]:
        run.add_break()
        run.text += line


def _set_run_fonts(run, *, bold: bool) -> None:
    run.font.name = LIBRE_BASKERVILLE
    run.font.bold = bold
    rfonts = run._element.get_or_add_rPr().get_or_add_rFonts()
    rfonts.set(qn("w:ascii"), LIBRE_BASKERVILLE)
    rfonts.set(qn("w:hAnsi"), LIBRE_BASKERVILLE)
    rfonts.set(qn("w:cs"), LIBRE_BASKERVILLE)
    rfonts.set(qn("w:eastAsia"), LIBRE_BASKERVILLE)


def _apply_body_paragraph(p, text: str) -> None:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(8)
    _add_multiline_run(p, text, bold=False, size_pt=12)


def _apply_main_heading_paragraph(p, text: str, *, page_break_before: bool) -> None:
    p.paragraph_format.page_break_before = page_break_before
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    p.paragraph_format.space_before = Pt(36)
    # Extra gap after title: default trailing space plus one line at 18pt × 1.2 line spacing.
    p.paragraph_format.space_after = Pt(24 + 22)
    _add_multiline_run(p, text, bold=True, size_pt=18)


def _apply_list_paragraph(p, text: str) -> None:
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.5
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(6)
    _add_multiline_run(p, text, bold=False, size_pt=12)


def _apply_sub_heading_paragraph(p, text: str) -> None:
    p.paragraph_format.page_break_before = False
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_before = Pt(16)
    p.paragraph_format.space_after = Pt(10)
    _add_multiline_run(p, text, bold=True, size_pt=14)


def _ensure_page_number_footer(doc: Document) -> None:
    for sec in doc.sections:
        footer = sec.footer
        while len(footer.paragraphs) > 1:
            p_el = footer.paragraphs[-1]._element
            p_el.getparent().remove(p_el)
        if footer.paragraphs:
            fp = footer.paragraphs[0]
            _clear_paragraph_runs(fp)
        else:
            fp = footer.add_paragraph()
        fp.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = fp.add_run()
        _set_run_fonts(run, bold=False)
        run.font.size = Pt(10)
        fld_begin = OxmlElement("w:fldChar")
        fld_begin.set(qn("w:fldCharType"), "begin")
        instr = OxmlElement("w:instrText")
        instr.set(qn("xml:space"), "preserve")
        instr.text = "PAGE"
        fld_end = OxmlElement("w:fldChar")
        fld_end.set(qn("w:fldCharType"), "end")
        run._r.append(fld_begin)
        run._r.append(instr)
        run._r.append(fld_end)


def _configure_document_defaults(doc: Document) -> None:
    for sec in doc.sections:
        sec.top_margin = Inches(1)
        sec.bottom_margin = Inches(1)
        sec.left_margin = Inches(1)
        sec.right_margin = Inches(1)

    normal = doc.styles["Normal"]
    normal.font.name = LIBRE_BASKERVILLE
    normal.font.size = Pt(12)
    normal.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    normal.paragraph_format.line_spacing = 1.5
    normal.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    normal.paragraph_format.space_after = Pt(8)
    nr = normal.element.get_or_add_rPr()
    nf = nr.get_or_add_rFonts()
    nf.set(qn("w:ascii"), LIBRE_BASKERVILLE)
    nf.set(qn("w:hAnsi"), LIBRE_BASKERVILLE)
    nf.set(qn("w:cs"), LIBRE_BASKERVILLE)
    nf.set(qn("w:eastAsia"), LIBRE_BASKERVILLE)

    _ensure_page_number_footer(doc)


def _style_table_cell(cell, text: str) -> None:
    t = strip_markdown_artifacts(text)
    cell.text = t
    for p in cell.paragraphs:
        _apply_body_paragraph(p, t)


def _ensure_translator_book_styles(doc: Document) -> None:
    for style_name in (
        STYLE_BOOK_TITLE,
        STYLE_BOOK_AUTHOR,
        STYLE_TOC_HEADING,
        STYLE_TOC_ENTRY,
    ):
        try:
            st = doc.styles.add_style(style_name, WD_STYLE_TYPE.PARAGRAPH)
        except ValueError:
            st = doc.styles[style_name]
        st.base_style = doc.styles["Normal"]
        st.font.name = LIBRE_BASKERVILLE
        st.font.bold = False
        st.font.size = Pt(12)
        nf = st.element.get_or_add_rPr().get_or_add_rFonts()
        nf.set(qn("w:ascii"), LIBRE_BASKERVILLE)
        nf.set(qn("w:hAnsi"), LIBRE_BASKERVILLE)
        nf.set(qn("w:cs"), LIBRE_BASKERVILLE)
        nf.set(qn("w:eastAsia"), LIBRE_BASKERVILLE)


def _title_font_size_pt(text: str) -> float:
    """Shrink cover title in Word when the string is very long (PDF uses tighter fitting)."""
    n = len((text or "").strip())
    if n <= 36:
        return 60.0
    if n <= 72:
        return 48.0
    if n <= 120:
        return 36.0
    if n <= 200:
        return 30.0
    return 26.0


def _toc_entry_indent_pt(block: ContentBlock, text: str) -> float:
    raw = text or ""
    lead = re.match(r"^(\s+)", raw)
    base = float(len(lead.group(1)) * 6) if lead else 0.0
    if block.type == BlockType.HEADING and block.level > 1:
        return min(72.0, 12.0 * (block.level - 1) + base)
    stripped = raw.lstrip()
    if re.match(r"^\d+(\.\d+)*\s", stripped):
        token = stripped.split()[0]
        depth = token.count(".")
        return min(72.0, base + 14.0 * max(0, depth))
    return min(48.0, base)


def _partition_export_blocks(
    blocks: list[ClassifiedBlock],
) -> tuple[
    list[ClassifiedBlock],
    list[ClassifiedBlock],
    list[ClassifiedBlock],
    list[ClassifiedBlock],
]:
    title: list[ClassifiedBlock] = []
    author: list[ClassifiedBlock] = []
    toc: list[ClassifiedBlock] = []
    body: list[ClassifiedBlock] = []
    for item in blocks:
        if item.action == SectionAction.OMIT:
            continue
        b = item.block
        if b.structural_tag == StructuralTag.TITLE and b.type != BlockType.TABLE:
            title.append(item)
        elif b.structural_tag == StructuralTag.AUTHOR and b.type != BlockType.TABLE:
            author.append(item)
        elif b.structural_tag == StructuralTag.TOC:
            toc.append(item)
        else:
            body.append(item)
    return title, author, toc, body


def _merge_title_text(items: list[ClassifiedBlock]) -> str:
    parts: list[str] = []
    for it in items:
        t = strip_markdown_artifacts(it.block.text or "").strip()
        if t:
            parts.append(t)
    return "\n".join(parts)


def _author_lines_from_items(items: list[ClassifiedBlock]) -> list[str]:
    lines: list[str] = []
    for it in items:
        t = strip_markdown_artifacts(it.block.text or "").strip()
        if not t:
            continue
        for part in t.split("\n"):
            s = part.strip()
            if s:
                lines.append(s)
    return lines


def _apply_book_title_paragraph(doc: Document, p, text: str) -> None:
    p.style = doc.styles[STYLE_BOOK_TITLE]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(24)
    p.paragraph_format.space_after = Pt(28)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.15
    _add_multiline_run(p, text, bold=True, size_pt=_title_font_size_pt(text))


def _apply_book_author_paragraph(doc: Document, p, text: str) -> None:
    p.style = doc.styles[STYLE_BOOK_AUTHOR]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    _add_multiline_run(p, text, bold=False, size_pt=24)


def _apply_toc_heading_paragraph(doc: Document, p, text: str) -> None:
    p.style = doc.styles[STYLE_TOC_HEADING]
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(16)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.2
    _add_multiline_run(p, text, bold=True, size_pt=16)


def _apply_toc_entry_paragraph(doc: Document, p, block: ContentBlock, text: str) -> None:
    p.style = doc.styles[STYLE_TOC_ENTRY]
    ind = _toc_entry_indent_pt(block, text)
    p.paragraph_format.left_indent = Pt(ind)
    p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.MULTIPLE
    p.paragraph_format.line_spacing = 1.25
    p.paragraph_format.space_after = Pt(4)
    p.paragraph_format.alignment = WD_ALIGN_PARAGRAPH.LEFT
    bold = bool(re.match(r"^\s*(chapter|part)\s+", text, re.I))
    _add_multiline_run(p, text, bold=bold, size_pt=12)


def _add_title_page_section(doc: Document, title: str, author_lines: list[str]) -> None:
    has_title = bool((title or "").strip())
    has_author = bool(author_lines)

    # No cover metadata: leave page 1 visually empty; main matter (chapters) starts page 2+.
    if not has_title and not has_author:
        spacer = doc.add_paragraph()
        spacer.paragraph_format.space_before = Pt(108)
        spacer.paragraph_format.space_after = Pt(0)
        _clear_paragraph_runs(spacer)
        doc.add_page_break()
        return

    spacer = doc.add_paragraph()
    spacer.paragraph_format.space_before = Pt(108)
    spacer.paragraph_format.space_after = Pt(0)
    _clear_paragraph_runs(spacer)

    if has_title:
        tp = doc.add_paragraph()
        _apply_book_title_paragraph(doc, tp, title.strip())

    if author_lines:
        for line in author_lines:
            ap = doc.add_paragraph()
            _apply_book_author_paragraph(doc, ap, line)
    elif has_title:
        pad = doc.add_paragraph()
        _apply_book_author_paragraph(doc, pad, "\u00a0")

    doc.add_page_break()


def _add_toc_section(doc: Document, toc_items: list[ClassifiedBlock]) -> None:
    for item in toc_items:
        b = item.block
        text = strip_markdown_artifacts(b.text or "")
        if b.type == BlockType.HEADING:
            p = doc.add_paragraph()
            _apply_toc_heading_paragraph(doc, p, text)
        elif b.type in (BlockType.PARAGRAPH, BlockType.LIST):
            p = doc.add_paragraph()
            _apply_toc_entry_paragraph(doc, p, b, text)
        elif b.type == BlockType.TABLE and b.data:
            rows = len(b.data)
            cols = max(len(r) for r in b.data) if b.data else 0
            if rows and cols:
                table = doc.add_table(rows=rows, cols=cols)
                table.style = "Table Grid"
                for ri, row in enumerate(b.data):
                    for ci in range(cols):
                        cell_text = row[ci] if ci < len(row) else ""
                        cell = table.rows[ri].cells[ci]
                        _style_table_cell(cell, cell_text)
                doc.add_paragraph()
    doc.add_page_break()


def _append_body_block(
    doc: Document, item: ClassifiedBlock, *, first_main_heading: list[bool]
) -> None:
    b = item.block
    raw = b.text or ""
    text = strip_markdown_artifacts(raw)

    if b.type == BlockType.HEADING:
        level = max(1, min(9, b.level))
        if level == 1:
            p = doc.add_paragraph(style="Heading 1")
            brk = not first_main_heading[0]
            first_main_heading[0] = False
            _apply_main_heading_paragraph(p, text, page_break_before=brk)
        else:
            p = doc.add_paragraph(style=f"Heading {level}")
            _apply_sub_heading_paragraph(p, text)
    elif b.type == BlockType.PARAGRAPH:
        p = doc.add_paragraph(style="Normal")
        _apply_body_paragraph(p, text)
    elif b.type == BlockType.LIST:
        p = doc.add_paragraph(style="List Bullet")
        _apply_list_paragraph(p, text)
    elif b.type == BlockType.TABLE and b.data:
        rows = len(b.data)
        cols = max(len(r) for r in b.data) if b.data else 0
        if rows and cols:
            table = doc.add_table(rows=rows, cols=cols)
            table.style = "Table Grid"
            for ri, row in enumerate(b.data):
                for ci in range(cols):
                    cell_text = row[ci] if ci < len(row) else ""
                    cell = table.rows[ri].cells[ci]
                    _style_table_cell(cell, cell_text)
            doc.add_paragraph()


def build_docx(blocks: list[ClassifiedBlock], output_path: Path) -> Path:
    doc = Document()
    _configure_document_defaults(doc)
    _ensure_translator_book_styles(doc)

    title_items, author_items, toc_items, body_items = _partition_export_blocks(blocks)
    _ = iter_auto_generated_toc_entries(body_items)

    title_text = _merge_title_text(title_items)
    author_lines = _author_lines_from_items(author_items)

    _add_title_page_section(doc, title_text, author_lines)

    if toc_items:
        _add_toc_section(doc, toc_items)

    first_flag = [True]
    for item in body_items:
        if item.action == SectionAction.OMIT:
            continue
        _append_body_block(doc, item, first_main_heading=first_flag)

    output_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(output_path))
    embed_libre_baskerville(output_path)
    return output_path.resolve()
