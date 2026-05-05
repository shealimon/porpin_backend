"""Build DOCX from :class:`~app.models.structured_document.StructuredDocument` only (no HTML).

Used by the translation export path so Word output mirrors the normalized structure
sidecar without reusing HTML/CSS.
"""

from __future__ import annotations

import re
from pathlib import Path

from docx import Document
from app.models.structured_document import (
    StructuredDocument,
    StructuredHeading,
    StructuredList,
    StructuredParagraph,
    StructuredTable,
)
from app.services.formatter.document_builder import (
    _add_title_page_section,
    _apply_toc_entry_paragraph,
    _apply_toc_heading_paragraph,
    _apply_body_paragraph,
    _apply_list_paragraph,
    _apply_main_heading_paragraph,
    _apply_sub_heading_paragraph,
    _configure_document_defaults,
    _ensure_translator_book_styles,
    _style_table_cell,
)
from app.services.formatter.docx_embed_fonts import embed_libre_baskerville
from app.services.formatter.chapter_heading_policy import (
    chapter_start_level,
    is_chapter_outline_level,
)
from app.models.document_models import ContentBlock, BlockType

_add_toc_entry_paragraph = _apply_toc_entry_paragraph


def _append_toc_node(doc: Document, node: StructuredHeading | StructuredParagraph | StructuredList | StructuredTable) -> None:
    if isinstance(node, StructuredHeading):
        p = doc.add_paragraph()
        _apply_toc_heading_paragraph(doc, p, node.text)
        return
    if isinstance(node, StructuredParagraph):
        p = doc.add_paragraph()
        fake = ContentBlock(type=BlockType.PARAGRAPH, text=node.text)
        _add_toc_entry_paragraph(doc, p, fake, node.text)
        return
    if isinstance(node, StructuredList):
        for it in node.items:
            p = doc.add_paragraph()
            fake = ContentBlock(type=BlockType.LIST, text=it)
            _add_toc_entry_paragraph(doc, p, fake, it)
        return
    if isinstance(node, StructuredTable):
        for row in node.rows:
            line = " | ".join(c.strip() for c in row)
            if line.strip():
                p = doc.add_paragraph()
                fake = ContentBlock(type=BlockType.PARAGRAPH, text=line)
                _add_toc_entry_paragraph(doc, p, fake, line)


def _add_toc_section_structured(doc: Document, toc_nodes: list) -> None:
    for node in toc_nodes:
        _append_toc_node(doc, node)
    doc.add_page_break()


def _append_body_node(
    doc: Document,
    node: StructuredHeading | StructuredParagraph | StructuredList | StructuredTable,
    *,
    body_started: list[bool],
    chapter_min_level: int | None,
) -> None:
    if isinstance(node, StructuredHeading):
        level = max(1, min(9, node.level))
        text = node.text
        if is_chapter_outline_level(level, chapter_min_level, heading_text=text):
            p = doc.add_paragraph(style="Heading 1")
            brk = body_started[0]
            _apply_main_heading_paragraph(p, text, page_break_before=brk)
        else:
            p = doc.add_paragraph(style=f"Heading {level}")
            _apply_sub_heading_paragraph(p, text)
        body_started[0] = True
        return
    if isinstance(node, StructuredParagraph):
        p = doc.add_paragraph(style="Normal")
        _apply_body_paragraph(p, node.text)
        body_started[0] = True
        return
    if isinstance(node, StructuredList):
        style_name = "List Number" if node.ordered else "List Bullet"
        for it in node.items:
            raw = (it or "").strip()
            if not raw:
                continue
            p = doc.add_paragraph(style=style_name)
            t = raw
            if node.ordered:
                t = re.sub(r"^\d+[\.)]\s+", "", t)
            else:
                t = re.sub(r"^[-*•]\s+", "", t)
            _apply_list_paragraph(p, t.strip())
        body_started[0] = True
        return
    if isinstance(node, StructuredTable):
        rows = node.rows
        if not rows:
            return
        nrows = len(rows)
        ncols = max(len(r) for r in rows) if rows else 0
        if not ncols:
            return
        table = doc.add_table(rows=nrows, cols=ncols)
        table.style = "Table Grid"
        for ri, row in enumerate(rows):
            for ci in range(ncols):
                cell_text = row[ci] if ci < len(row) else ""
                cell = table.rows[ri].cells[ci]
                _style_table_cell(cell, cell_text)
        doc.add_paragraph()

    body_started[0] = True


def build_docx_from_structured(doc: StructuredDocument, output_path: Path) -> Path:
    """Write a new DOCX from plain structure (title, authors, body, optional TOC)."""
    document = Document()
    _configure_document_defaults(document)
    _ensure_translator_book_styles(document)

    title_text = (doc.title or "").strip()
    author_lines = [a.strip() for a in doc.authors if (a or "").strip()]
    _add_title_page_section(document, title_text, author_lines)

    toc_nodes = [b for b in doc.content if b.content_tag == "toc"]
    body_nodes = [b for b in doc.content if b.content_tag != "toc"]

    if toc_nodes:
        _add_toc_section_structured(document, toc_nodes)

    ol_lvls = [
        n.level for n in body_nodes if isinstance(n, StructuredHeading)
    ]
    chapter_min = chapter_start_level(ol_lvls)

    body_started = [False]
    for node in body_nodes:
        _append_body_node(
            document,
            node,
            body_started=body_started,
            chapter_min_level=chapter_min,
        )

    output_path.parent.mkdir(parents=True, exist_ok=True)
    document.save(str(output_path))
    embed_libre_baskerville(output_path)
    return output_path.resolve()
